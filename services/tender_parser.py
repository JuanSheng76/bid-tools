"""招标文件解析服务

解析 .docx / .pdf 招标文件，提取：
- 资格要求（密封/盖章/递交/证书/承诺函/保证金）
- 评分标准（评分项/分值/评分方法）
- 重要注意事项（映射到 task_type）
- 智能推荐（匹配公司资质/业绩/人员）

纯规则匹配：章节标题关键词定位 + 表格解析 + 正则提取

性能说明：
- 所有正则模式在模块加载时预编译，避免每次请求重复编译
- extract_important_notes 使用单次扫描替代多次逐关键词扫描
- 资格提取合并为少量 pass，避免重复遍历
"""
import os
import re
import uuid
from datetime import datetime
from typing import Optional

from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.orm import selectinload

from config import BASE_DIR

# ====== 预编译正则（模块级，只编译一次） ======

# 评分项分值提取
_SCORING_PATTERNS = [
    re.compile(r'([^。；\n,，]{2,30}?)[（(]\s*(\d+)\s*分[)）]'),
    re.compile(r'([^。；\n,，]{2,30}?)[：:]\s*(\d+)\s*分'),
    re.compile(r'([^。；\n,，]{2,30}?)\s+(\d+)\s*分'),
]

# 金额提取
_RE_BOND_AMOUNT_WAN = re.compile(r'(?:投标保证金|保证金).*?([\d,.]+)\s*万')
_RE_BOND_AMOUNT_YUAN = re.compile(r'(?:投标保证金|保证金).*?([\d,]+)\s*元')
_RE_PERF_BOND = re.compile(r'履约保证金.*?([\d,.]+)\s*万')
_RE_COPIES = re.compile(r'([\d]+)\s*份')
_RE_COPIES_ZB = re.compile(r'正本\s*([\d]+)\s*份.*?副本\s*([\d]+)\s*份')
_RE_BUDGET_AMOUNT = re.compile(r'(\d+)\s*万')
_RE_RECENT_YEARS = re.compile(r'近\s*(\d+)\s*年')
_RE_DEADLINE_DATE = re.compile(r'(\d{4}[-/年]\d{1,2}[-/月]\d{1,2}[日]?\s*\d{1,2}:\d{2})')
_RE_LOCATION = re.compile(r'(?:递交|送达|开标)\s*地点[：:]\s*(.+?)(?:[。；\n]|$)')

# 注意事项预编译：构建 (keyword_pattern, task_type, priority) 列表
# 每个条目是一个预编译的正则，匹配 "包含关键词的句子"
_NOTE_PATTERNS: list[tuple[re.Pattern, str, str]] = []
_NOTE_KEYWORD_TO_META: dict[str, tuple[str, str]] = {}  # keyword -> (task_type, priority)

_NOTE_TASK_TYPE_MAP_RAW = [
    (["盖章", "签字", "签章", "逐页", "骑缝"], "stamp", "urgent"),
    (["份数", "正本", "副本", "密封", "装订", "封装"], "format", "high"),
    (["保证金", "保函", "到账", "标书费", "电汇"], "pricing", "high"),
    (["证书复印件", "业绩证明", "验收报告", "社保证明", "原件", "复印件"], "certs", "high"),
    (["承诺函格式", "声明函", "法人证明", "授权委托书", "信用中国", "无行贿"], "qualifications", "high"),
    (["质疑期限", "答疑", "踏勘", "澄清", "答疑会"], "get_docs", "medium"),
]

for _keywords, _task_type, _priority in _NOTE_TASK_TYPE_MAP_RAW:
    for _kw in _keywords:
        # 预编译：匹配包含关键词的完整句子（以。；\n为界）
        _NOTE_PATTERNS.append((
            re.compile(r'[^。；\n]*' + re.escape(_kw) + r'[^。；\n]*'),
            _task_type,
            _priority,
        ))
        _NOTE_KEYWORD_TO_META[_kw] = (_task_type, _priority)

# 证书/承诺函关键词合并为单次扫描模式
_CERT_PATTERN = re.compile('|'.join(re.escape(k) for k in [
    "营业执照", "资质证书", "许可证", "认证证书", "ISO",
    "检验报告", "检测报告", "安全生产许可证", "建筑业企业资质",
    "承装.*修.*试", "安全生产考核", "特种作业", "职业健康",
    "质量管理体系", "环境管理体系", "3C认证", "CCC",
    "计量认证", "CMA", "CNAS",
]))

_COMMITMENT_PATTERN = re.compile('|'.join(re.escape(k) for k in [
    "承诺函", "声明函", "承诺书", "声明", "证明",
    "无行贿犯罪", "无重大违法", "诚信", "信用中国",
    "无失信", "非联合体", "不转包", "不分包",
    "授权委托书", "法定代表人", "法人授权",
]))

# 密封/盖章/递交关键词集合
_SEALING_KW_SET = frozenset(["密封", "封套", "密封袋", "密封条", "封条", "正本", "副本", "份数", "一式"])
_STAMPING_KW_SET = frozenset(["盖章", "签字", "签章", "逐页", "骑缝", "加盖公章", "法定代表人或其委托代理人", "签字或盖章"])
_SUBMISSION_KW_SET = frozenset(["递交截止", "投标截止", "递交地点", "递交方式", "邮寄", "送达", "递交时间"])

# ====== 章节关键词映射 ======

SECTION_KEYWORDS = {
    "bidder_instructions": [
        "投标人须知", "投标须知", "供应商须知", "须知前附表",
        "投标人须知前附表", "前附表",
    ],
    "qualification": [
        "资格要求", "投标人资格", "资格条件", "资格审查",
        "合格投标人", "投标人资质", "资质要求", "资质条件",
    ],
    "scoring": [
        "评分办法", "评分标准", "评标办法", "评审办法",
        "评审标准", "综合评分", "评分细则", "详细评审",
        "评标方法", "评审方法",
    ],
    "bid_doc_format": [
        "投标文件格式", "投标文件组成", "投标文件编制",
        "响应文件格式", "投标文件内容", "投标文件要求",
    ],
    "project_requirements": [
        "项目需求", "技术规格", "服务要求", "采购需求",
        "招标范围", "项目概况", "技术要求",
    ],
}


# ====== .docx 解析 ======

def _extract_docx_text(doc) -> list[dict]:
    """从 python-docx Document 提取段落列表，包含样式信息"""
    paragraphs = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        style_name = para.style.name if para.style else ""
        # 判断是否为标题样式
        is_heading = "Heading" in style_name or "heading" in style_name or "标题" in style_name
        # 获取字体信息判断标题
        bold = False
        font_size = None
        for run in para.runs:
            if run.bold:
                bold = True
            if run.font.size:
                font_size = run.font.size.pt
                break
        paragraphs.append({
            "index": i,
            "text": text,
            "style": style_name,
            "is_heading": is_heading,
            "bold": bold,
            "font_size": font_size,
        })
    return paragraphs


def _extract_docx_tables(doc) -> list[list[list[str]]]:
    """提取所有表格为二维字符串数组"""
    tables = []
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(cells)
        tables.append(rows)
    return tables


def _locate_sections(paragraphs: list[dict]) -> dict:
    """定位章节边界，返回 {section_key: (start_index, end_index)}"""
    sections = {}
    current_section = None
    current_start = 0

    for i, p in enumerate(paragraphs):
        matched = None
        for section_key, keywords in SECTION_KEYWORDS.items():
            if _is_section_heading(p, keywords):
                matched = section_key
                break

        if matched:
            # 保存上一个章节
            if current_section and current_section not in sections:
                sections[current_section] = (current_start, i)
            current_section = matched
            current_start = i

    # 最后一个章节
    if current_section and current_section not in sections:
        sections[current_section] = (current_start, len(paragraphs))

    return sections


def _is_section_heading(para: dict, keywords: list[str]) -> bool:
    """判断段落是否为指定章节的标题"""
    text = para["text"]
    # 章节标题通常较短（< 30 字）
    if len(text) > 50:
        return False

    for kw in keywords:
        if kw in text:
            # 标题样式 或 加粗 或 字号较大 或 编号开头
            if para["is_heading"] or para["bold"]:
                return True
            if para["font_size"] and para["font_size"] >= 13:
                return True
            # 中文编号开头：第x章、一、二、1.、1、
            if re.match(r'^[（(]?[一二三四五六七八九十\d]+[)）．、.]', text):
                return True
            if re.match(r'^第[一二三四五六七八九十\d]+[章节条]', text):
                return True
    return False


def _get_section_text(paragraphs: list[dict], sections: dict, key: str) -> str:
    """获取某个章节的全部文本"""
    if key not in sections:
        return ""
    start, end = sections[key]
    return "\n".join(p["text"] for p in paragraphs[start:end])


# ====== 资格要求提取 ======

def _extract_certificates_from_text(text: str) -> list[dict]:
    """从文本中提取所需证书列表（单次扫描）"""
    certs = []
    seen = set()
    for match in _CERT_PATTERN.finditer(text):
        m = match.group()
        if m not in seen:
            seen.add(m)
            certs.append({"name": m, "type": "basic", "detail": "", "matched": False})
    return certs


def _extract_commitments_from_text(text: str) -> list[dict]:
    """从文本中提取所需承诺函列表（单次扫描）"""
    commitments = []
    seen = set()
    for match in _COMMITMENT_PATTERN.finditer(text):
        m = match.group()
        if m not in seen:
            seen.add(m)
            commitments.append({"name": m, "detail": ""})
    return commitments


def _extract_sealing_requirements(text: str) -> dict:
    """提取密封要求"""
    result = {"copies": "", "packaging": "", "requirements_text": ""}
    copies_match = _RE_COPIES.search(text)
    if copies_match:
        result["copies"] = copies_match.group(0)
    zb_match = _RE_COPIES_ZB.search(text)
    if zb_match:
        result["copies"] = f"正本{zb_match.group(1)}份，副本{zb_match.group(2)}份"

    seals = [kw for kw in _SEALING_KW_SET if kw in text]
    result["packaging"] = "、".join(seals) if seals else ""
    result["requirements_text"] = _extract_surrounding_text(text, _SEALING_KW_SET)
    return result


def _extract_stamping_requirements(text: str) -> dict:
    """提取盖章要求"""
    stamps = [kw for kw in _STAMPING_KW_SET if kw in text]
    return {
        "requirements": stamps,
        "requirements_text": _extract_surrounding_text(text, _STAMPING_KW_SET),
    }


def _extract_submission_requirements(text: str) -> dict:
    """提取递交要求"""
    result = {"deadline": "", "location": "", "method": "", "requirements_text": ""}
    deadline_match = _RE_DEADLINE_DATE.search(text)
    if deadline_match:
        result["deadline"] = deadline_match.group(1)
    location_match = _RE_LOCATION.search(text)
    if location_match:
        result["location"] = location_match.group(1).strip()
    if "邮寄" in text:
        result["method"] = "邮寄"
    elif "现场递交" in text or "现场提交" in text:
        result["method"] = "现场递交"
    result["requirements_text"] = _extract_surrounding_text(text, _SUBMISSION_KW_SET)
    return result


def _extract_financial_requirements(text: str) -> dict:
    """提取保证金等财务要求"""
    result = {"bid_bond": None, "bid_bond_form": "", "performance_bond": None, "other": ""}
    bond_match = _RE_BOND_AMOUNT_WAN.search(text)
    if not bond_match:
        bond_match = _RE_BOND_AMOUNT_YUAN.search(text)
        if bond_match:
            try:
                result["bid_bond"] = float(bond_match.group(1).replace(",", "")) / 10000
            except ValueError:
                pass
    else:
        try:
            result["bid_bond"] = float(bond_match.group(1).replace(",", ""))
        except ValueError:
            pass

    if "保函" in text:
        result["bid_bond_form"] = "保函"
    elif "电汇" in text or "转账" in text:
        result["bid_bond_form"] = "电汇/转账"

    perf_match = _RE_PERF_BOND.search(text)
    if perf_match:
        try:
            result["performance_bond"] = float(perf_match.group(1).replace(",", ""))
        except ValueError:
            pass

    return result


def _extract_surrounding_text(text: str, keywords: frozenset) -> str:
    """提取包含关键字的上下文句子（单次 split + 过滤）"""
    sentences = re.split(r'[。；\n]', text)
    relevant = [s.strip() for s in sentences if any(kw in s for kw in keywords)]
    return "\n".join(relevant[:10])


def _extract_qualification_requirements(paragraphs: list[dict], sections: dict, full_text: str) -> dict:
    """主入口：提取资格要求"""
    qual_text = _get_section_text(paragraphs, sections, "qualification")
    bidder_text = _get_section_text(paragraphs, sections, "bidder_instructions")
    # 只取前 8000 字符做详细扫描（资格要求在文档前半部，扫描全文浪费且多余）
    combined = f"{bidder_text}\n{qual_text}\n{full_text[:8000]}"

    return {
        "sealing": _extract_sealing_requirements(combined),
        "submission": _extract_submission_requirements(combined),
        "stamping": _extract_stamping_requirements(combined),
        "required_certificates": _extract_certificates_from_text(combined),
        "required_commitments": _extract_commitments_from_text(combined),
        "financial_requirements": _extract_financial_requirements(combined),
        "raw_text": qual_text or combined[:3000],
    }


# ====== 评分标准提取 ======

_SCORING_CATEGORY_MAP = [
    (["报价", "价格", "投标报价", "投标价格"], "price"),
    (["技术", "技术方案", "技术部分", "服务方案", "实施方案", "施工组织"], "technical"),
    (["业绩", "项目业绩", "类似项目", "过往业绩", "合同业绩"], "performance"),
    (["人员", "项目负责人", "团队", "项目经理", "技术人员", "人员配置"], "personnel"),
    (["资质", "资格", "企业资质", "资质等级", "信誉"], "qualification"),
]


def _infer_scoring_category(label: str) -> str:
    """根据评分项名称推断类别"""
    for keywords, category in _SCORING_CATEGORY_MAP:
        for kw in keywords:
            if kw in label:
                return category
    return "other"


def _parse_scoring_from_text(text: str) -> list[dict]:
    """从文本中提取评分项（使用预编译正则）"""
    items = []
    for pattern in _SCORING_PATTERNS:
        for label, points in pattern.findall(text):
            label = label.strip()
            if label and len(label) >= 2:
                items.append({
                    "category": _infer_scoring_category(label),
                    "label": label,
                    "max_points": int(points),
                    "scoring_method": "",
                    "requirements": [],
                })

    # 去重（按 label 去重，保留分值更高的）
    seen = {}
    for item in items:
        key = item["label"]
        if key not in seen or item["max_points"] > seen[key]["max_points"]:
            seen[key] = item
    return list(seen.values())


def _parse_scoring_table_rows(rows: list[list[str]]) -> list[dict]:
    """从表格行中解析评分项（常见格式：序号|评分项|分值|评分标准）"""
    items = []
    for row in rows:
        if len(row) < 2:
            continue
        # 跳过表头
        if any(h in "".join(row) for h in ["序号", "评分项", "项目", "类别"]):
            continue

        # 在整行中找分值
        text = " ".join(row)
        points_match = re.search(r'(\d+)\s*分', text)
        if not points_match:
            continue

        points = int(points_match.group(1))
        # 找最像评分项名称的单元格（排除数字、纯分值）
        label = ""
        for cell in row:
            cell = cell.strip()
            if len(cell) >= 2 and not re.match(r'^[\d.\s]+$', cell) and "分" not in cell:
                label = cell
                break

        if label and points:
            items.append({
                "category": _infer_scoring_category(label),
                "label": label,
                "max_points": points,
                "scoring_method": text,
                "requirements": [],
            })

    return items


def _extract_scoring_criteria(paragraphs: list[dict], sections: dict, tables: list[list[list[str]]], full_text: str) -> dict:
    """主入口：提取评分标准"""
    scoring_text = _get_section_text(paragraphs, sections, "scoring")

    # 先从表格中提取
    items = []
    for table in tables:
        table_items = _parse_scoring_table_rows(table)
        if table_items:
            items.extend(table_items)

    # 再从文本中补充
    if scoring_text:
        text_items = _parse_scoring_from_text(scoring_text)
        items.extend(text_items)

    # 如果评分章节没有，从全文提取
    if not items:
        items = _parse_scoring_from_text(full_text)

    # 去重 + 计算总分
    seen = {}
    for item in items:
        key = item["label"]
        if key not in seen or item["max_points"] > seen[key]["max_points"]:
            seen[key] = item
    items = list(seen.values())

    total_points = sum(i["max_points"] for i in items) or 100

    return {
        "total_points": total_points,
        "items": items,
        "raw_text": scoring_text or full_text[:3000],
    }


# ====== 注意事项提取 ======

def extract_important_notes(full_text: str) -> list[dict]:
    """单次扫描提取重要注意事项并映射到 task_type

    优化：预编译所有关键词正则为 module-level _NOTE_PATTERNS，
    单次遍历 full_text，按句子匹配，每个句子只检查一次。
    """
    notes = []
    seen_texts = set()

    # 按。；\n 拆句，逐句匹配（比全文 finditer 更可控）
    sentences = re.split(r'[。；\n]+', full_text)
    for s in sentences:
        s = s.strip()
        if len(s) < 4:
            continue
        for pattern, task_type, priority in _NOTE_PATTERNS:
            if pattern.search(s):
                if s not in seen_texts:
                    seen_texts.add(s)
                    notes.append({
                        "text": s,
                        "task_type": task_type,
                        "priority": priority,
                    })
                break  # 一句只归入一个类型

    # 按优先级排序
    priority_order = {"urgent": 0, "high": 1, "medium": 2, "low": 3}
    notes.sort(key=lambda n: priority_order.get(n["priority"], 99))

    return notes[:30]


# ====== 匹配推荐引擎 ======

def match_qualifications(scoring: dict, company_qualifications: list) -> list[dict]:
    """匹配公司资质与招标要求"""
    if not company_qualifications:
        return []

    # 提取评分中的资质关键词
    qual_keywords = set()
    scoring_text = scoring.get("raw_text", "")
    for item in scoring.get("items", []):
        scoring_text += item.get("label", "") + " " + item.get("scoring_method", "")

    for kw in ["资质", "证书", "许可证", "认证", "ISO", "等级"]:
        if kw in scoring_text:
            qual_keywords.add(kw)

    results = []
    for idx, q in enumerate(company_qualifications):
        qual_text = f"{q.get('name', '')} {q.get('level', '')} {q.get('issuing_authority', '')}"
        score = 0.5  # 基础分

        # 关键词匹配加分
        for kw in qual_keywords:
            if kw in qual_text:
                score += 0.3

        # 有证书编号 + 有等级 加分
        if q.get("cert_no"):
            score += 0.1
        if q.get("level"):
            score += 0.1

        results.append({
            "from_db_index": idx,
            "name": q.get("name", ""),
            "level": q.get("level", ""),
            "match_score": min(round(score, 2), 1.0),
        })

    results.sort(key=lambda r: r["match_score"], reverse=True)
    return results


def match_performances(scoring: dict, company_performances: list) -> list[dict]:
    """匹配公司业绩与招标要求"""
    if not company_performances:
        return []

    scoring_text = scoring.get("raw_text", "")
    # 提取业绩要求：金额/时间/类型
    amount_req = None
    amount_match = _RE_BUDGET_AMOUNT.search(scoring_text)
    if amount_match:
        amount_req = float(amount_match.group(1))

    year_req = None
    year_match = _RE_RECENT_YEARS.search(scoring_text)
    if year_match:
        year_req = int(year_match.group(1))

    project_types = []
    for kw in ["光伏", "检测", "检验", "工程", "服务", "运维", "采购"]:
        if kw in scoring_text:
            project_types.append(kw)

    results = []
    for idx, p in enumerate(company_performances):
        score = 0.3  # 基础分
        p_text = f"{p.get('project_name', '')} {p.get('project_type', '')} {p.get('description', '')}"

        # 项目类型匹配
        for pt in project_types:
            if pt in p_text:
                score += 0.5

        # 金额匹配
        contract_amount = p.get("contract_amount")
        if amount_req and contract_amount:
            if contract_amount >= amount_req:
                score += 0.3
            elif contract_amount >= amount_req * 0.5:
                score += 0.1

        # 时间匹配
        contract_date = p.get("contract_date", "")
        if year_req and contract_date:
            try:
                year = int(contract_date[:4])
                current_year = datetime.utcnow().year
                if current_year - year <= year_req:
                    score += 0.2
            except (ValueError, IndexError):
                pass

        # 描述完整度
        if p.get("description") and len(p.get("description", "")) > 50:
            score += 0.1
        if p.get("client_name"):
            score += 0.05

        results.append({
            "from_db_index": idx,
            "project_name": p.get("project_name", ""),
            "match_score": min(round(score, 2), 1.0),
            "contract_amount": contract_amount,
            "contract_date": contract_date,
            "client_name": p.get("client_name", ""),
        })

    results.sort(key=lambda r: (r["match_score"], -(r["contract_amount"] or 0)), reverse=True)
    return results


def match_personnel(scoring: dict, company_personnel: list) -> list[dict]:
    """匹配公司人员与招标要求"""
    if not company_personnel:
        return []

    scoring_text = scoring.get("raw_text", "")
    # 提取人员证书关键词
    cert_keywords = set()
    for kw in ["建造师", "工程师", "安全员", "项目经理", "检测", "光伏", "电气"]:
        if kw in scoring_text:
            cert_keywords.add(kw)

    results = []
    for idx, p in enumerate(company_personnel):
        score = 0.3
        p_text = f"{p.get('name', '')} {p.get('position', '')} {p.get('certifications', '')}"

        # 证书关键词匹配
        for kw in cert_keywords:
            if kw in p_text:
                score += 0.4

        # 有证书
        if p.get("certifications"):
            score += 0.2

        # 有职位
        if p.get("position"):
            score += 0.1

        results.append({
            "from_db_index": idx,
            "name": p.get("name", ""),
            "match_score": min(round(score, 2), 1.0),
            "position": p.get("position", ""),
            "certifications": p.get("certifications", ""),
        })

    results.sort(key=lambda r: r["match_score"], reverse=True)
    return results


# ====== 任务清单同步 ======

async def enrich_task_checklists(notice_id: str, db: AsyncSession, important_notes: list[dict]) -> int:
    """将注意事项追加到对应 task_type 的 Task.checklist 中（去重）"""
    from models import Task

    result = await db.execute(
        select(Task).where(Task.notice_id == notice_id)
    )
    tasks = result.scalars().all()

    if not tasks or not important_notes:
        return 0

    # 按 task_type 分组现有任务
    tasks_by_type = {}
    for t in tasks:
        tasks_by_type.setdefault(t.task_type, []).append(t)

    added_count = 0

    for note in important_notes:
        task_type = note["task_type"]
        if task_type not in tasks_by_type:
            continue

        for task in tasks_by_type[task_type]:
            checklist = list(task.checklist or [])
            existing_texts = {item.get("text", "") for item in checklist}

            note_text = f"[招标文件] {note['text']}"
            if note_text not in existing_texts:
                checklist.append({"text": note_text, "done": False})
                task.checklist = checklist
                added_count += 1
                break  # 每个 note 只追加到一个 task

    if added_count > 0:
        await db.commit()

    return added_count


# ====== 主入口 ======

def parse_tender_docx(file_path: str, original_filename: str, company=None) -> dict:
    """解析 .docx 招标文件，返回完整 tender_analysis dict"""
    from docx import Document as DocxDocument

    print(f"[parser:docx] 开始加载文档: {file_path}", flush=True)
    doc = DocxDocument(file_path)
    print(f"[parser:docx] 文档加载完成, 段落数={len(doc.paragraphs)}, 表格数={len(doc.tables)}", flush=True)

    try:
        paragraphs = _extract_docx_text(doc)
        print(f"[parser:docx] 段落提取完成, 有效段落={len(paragraphs)}", flush=True)
    except Exception as e:
        raise RuntimeError(f"段落提取失败: {e}") from e

    try:
        tables = _extract_docx_tables(doc)
        print(f"[parser:docx] 表格提取完成, 表格数={len(tables)}", flush=True)
    except Exception as e:
        raise RuntimeError(f"表格提取失败: {e}") from e

    full_text = "\n".join(p["text"] for p in paragraphs)

    try:
        sections = _locate_sections(paragraphs)
        print(f"[parser:docx] 章节定位完成, 找到章节={list(sections.keys())}", flush=True)
    except Exception as e:
        raise RuntimeError(f"章节定位失败: {e}") from e

    try:
        qual_reqs = _extract_qualification_requirements(paragraphs, sections, full_text)
        print(f"[parser:docx] 资格要求提取完成", flush=True)
    except Exception as e:
        raise RuntimeError(f"资格要求提取失败: {e}") from e

    try:
        scoring = _extract_scoring_criteria(paragraphs, sections, tables, full_text)
        print(f"[parser:docx] 评分标准提取完成, 评分项={len(scoring.get('items',[]))}", flush=True)
    except Exception as e:
        raise RuntimeError(f"评分标准提取失败: {e}") from e

    try:
        important_notes = extract_important_notes(full_text)
        print(f"[parser:docx] 注意事项提取完成, 条数={len(important_notes)}", flush=True)
    except Exception as e:
        raise RuntimeError(f"注意事项提取失败: {e}") from e

    try:
        recommendations = {"qualifications": [], "performances": [], "personnel": [], "generated_at": datetime.utcnow().isoformat()}
        if company:
            recommendations["qualifications"] = match_qualifications(scoring, company.qualifications or [])
            recommendations["performances"] = match_performances(scoring, company.performances or [])
            recommendations["personnel"] = match_personnel(scoring, company.personnel or [])
        print(f"[parser:docx] 匹配推荐完成", flush=True)
    except Exception as e:
        raise RuntimeError(f"匹配推荐失败: {e}") from e

    try:
        for cert in qual_reqs.get("required_certificates", []):
            if company:
                cert_name = cert["name"]
                for q in (company.qualifications or []):
                    if cert_name in q.get("name", "") or cert_name in q.get("level", ""):
                        cert["matched"] = True
                        break
        print(f"[parser:docx] 证书匹配完成", flush=True)
    except Exception as e:
        raise RuntimeError(f"证书匹配失败: {e}") from e

    return {
        "file_name": original_filename,
        "file_stored_at": file_path,
        "parsed_at": datetime.utcnow().isoformat(),
        "parse_version": 1,
        "qualification_requirements": qual_reqs,
        "scoring_criteria": scoring,
        "recommendations": recommendations,
        "important_notes": important_notes,
    }


def parse_tender_pdf(file_path: str, original_filename: str, company=None) -> dict:
    """解析 .pdf 招标文件

    策略：pdfplumber 提取所有文本 + 表格 → 用同样的正则/关键词规则解析
    pdfplumber 无 Heading 样式概念，依赖字号+加粗判断章节标题
    """
    import pdfplumber

    all_text_lines = []
    paragraphs = []
    tables = []

    print(f"[parser:pdf] 开始加载 PDF: {file_path}", flush=True)
    with pdfplumber.open(file_path) as pdf:
        total_pages = len(pdf.pages)
        print(f"[parser:pdf] PDF 加载完成, 总页数={total_pages}", flush=True)
        for page_num, page in enumerate(pdf.pages):
            try:
                text = page.extract_text()
                if text:
                    for line in text.split("\n"):
                        line = line.strip()
                        if line:
                            all_text_lines.append(line)

                page_tables = page.extract_tables()
                for t in page_tables:
                    if t:
                        tables.append(t)
            except Exception as e:
                print(f"[parser:pdf] 第{page_num+1}页处理异常: {e}", flush=True)
                # 单页异常不中断整体解析
            if (page_num + 1) % 10 == 0:
                print(f"[parser:pdf] 已处理 {page_num+1}/{total_pages} 页", flush=True)

    print(f"[parser:pdf] 文本提取完成, 行数={len(all_text_lines)}, 表格数={len(tables)}", flush=True)

    # 将 pdf 文本行转换为 paragraphs 格式（模拟 docx 结构）
    for i, line in enumerate(all_text_lines):
        is_heading = False
        if len(line) < 30:
            for keywords in SECTION_KEYWORDS.values():
                for kw in keywords:
                    if kw in line:
                        is_heading = True
                        break
                if is_heading:
                    break
        if re.match(r'^[（(]?[一二三四五六七八九十\d]+[)）．、.]', line):
            is_heading = True
        paragraphs.append({
            "index": i, "text": line, "style": "",
            "is_heading": is_heading, "bold": False, "font_size": None,
        })

    full_text = "\n".join(all_text_lines)

    try:
        sections = _locate_sections(paragraphs)
        print(f"[parser:pdf] 章节定位完成, 找到={list(sections.keys())}", flush=True)
    except Exception as e:
        raise RuntimeError(f"章节定位失败: {e}") from e

    try:
        qual_reqs = _extract_qualification_requirements(paragraphs, sections, full_text)
        print(f"[parser:pdf] 资格要求提取完成", flush=True)
    except Exception as e:
        raise RuntimeError(f"资格要求提取失败: {e}") from e

    try:
        scoring = _extract_scoring_criteria(paragraphs, sections, tables, full_text)
        print(f"[parser:pdf] 评分标准提取完成, 评分项={len(scoring.get('items',[]))}", flush=True)
    except Exception as e:
        raise RuntimeError(f"评分标准提取失败: {e}") from e

    try:
        important_notes = extract_important_notes(full_text)
        print(f"[parser:pdf] 注意事项提取完成, 条数={len(important_notes)}", flush=True)
    except Exception as e:
        raise RuntimeError(f"注意事项提取失败: {e}") from e

    try:
        recommendations = {"qualifications": [], "performances": [], "personnel": [], "generated_at": datetime.utcnow().isoformat()}
        if company:
            recommendations["qualifications"] = match_qualifications(scoring, company.qualifications or [])
            recommendations["performances"] = match_performances(scoring, company.performances or [])
            recommendations["personnel"] = match_personnel(scoring, company.personnel or [])
        print(f"[parser:pdf] 匹配推荐完成", flush=True)
    except Exception as e:
        raise RuntimeError(f"匹配推荐失败: {e}") from e

    try:
        for cert in qual_reqs.get("required_certificates", []):
            if company:
                cert_name = cert["name"]
                for q in (company.qualifications or []):
                    if cert_name in q.get("name", "") or cert_name in q.get("level", ""):
                        cert["matched"] = True
                        break
        print(f"[parser:pdf] 证书匹配完成", flush=True)
    except Exception as e:
        raise RuntimeError(f"证书匹配失败: {e}") from e

    return {
        "file_name": original_filename,
        "file_stored_at": file_path,
        "parsed_at": datetime.utcnow().isoformat(),
        "parse_version": 1,
        "qualification_requirements": qual_reqs,
        "scoring_criteria": scoring,
        "recommendations": recommendations,
        "important_notes": important_notes,
    }
